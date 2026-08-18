"""Regression tests for the parsing robustness fixes:

- DOCX tables are now included (skills are commonly stored in tables).
- NUL bytes / blank-line runs are normalized out of extracted text.
- The OCR fallback degrades gracefully (returns "") when OCR deps or
  pypdfium2 are missing.
"""
import io

import pytest

from app.services.resume_parser import parse_docx, _normalize_text, _ocr_fallback


def _docx_with_table() -> bytes:
    """In-memory DOCX containing a paragraph + a skills table."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Software Engineer")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "SQL"
    table.cell(1, 0).text = "Docker"
    table.cell(1, 1).text = "Kubernetes"
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_docx_table_cells_are_extracted():
    try:
        import docx  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")
    text = parse_docx(_docx_with_table())
    assert "Python" in text
    assert "SQL" in text
    assert "Docker" in text
    assert "Kubernetes" in text


def test_normalize_text_strips_nul_and_blank_runs():
    raw = "\x00\x00Python\x00\x00\n\n\nSQL\n\n\n\nDocker"
    assert _normalize_text(raw) == "Python\n\nSQL\n\nDocker"


def test_ocr_fallback_graceful_when_deps_missing(monkeypatch):
    """With pypdfium2 unavailable (e.g. fresh minimal env), the fallback must
    return '' — never raise — so the caller surfaces the standard no-text error."""
    called_with = {}

    def fake_import(name, *args, **kwargs):
        if name == "pytesseract":
            called_with["pytesseract"] = True
            raise ImportError("mock: pytesseract unavailable")
        return __import__(name, *args, **kwargs)

    monkeypatch.setattr("builtins.__import__", fake_import)
    assert _ocr_fallback(b"not-a-pdf") == ""
    assert called_with["pytesseract"] is True
