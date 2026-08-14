"""Parses resume files (PDF/DOCX) into plain text, in-memory only.

Privacy contract: no file is ever written to disk. Everything operates on the
UploadFile's in-memory buffer; nothing outlives the request.

Error contract: unsupported/corrupt input raises ValueError with a
user-presentable message — API layer maps that to HTTP 400, never a raw 500.
"""
import io

from app.utils.logging import get_logger

logger = get_logger(__name__)

MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB upload cap


def parse_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text_chunks: list[str] = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
        text = _normalize_text("\n".join(text_chunks))
    except Exception as exc:
        logger.warning("pdf_parse_failed", error=str(exc))
        raise ValueError(
            "Could not read this PDF. Try a text-based PDF or paste the text instead."
        ) from exc

    if not text:
        text = _ocr_fallback(file_bytes)
    if not text:
        raise ValueError(
            "No extractable text found. This PDF may be scanned — try the OCR option or paste the text."
        )
    return text


def _normalize_text(text: str) -> str:
    """Collapses runs of blank lines and strips NUL bytes that some writers
    (e.g. pdfplumber on certain generators) emit in extracted text."""
    text = text.replace("\x00", "")
    lines = [ln.strip() for ln in text.splitlines()]
    cleaned: list[str] = []
    for ln in lines:
        if ln:
            cleaned.append(ln)
        elif cleaned and cleaned[-1] != "":  # single blank separator between blocks
            cleaned.append("")
    return "\n".join(cleaned).strip()


def parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text.strip() for p in doc.paragraphs if p and p.text.strip()]
        # Skills are frequently placed in tables; the paragraph pass alone
        # would silently drop them.
        seen: set[str] = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text.lower() not in seen:
                        seen.add(cell_text.lower())
                        parts.append(cell_text)
    except Exception as exc:
        logger.warning("docx_parse_failed", error=str(exc))
        raise ValueError(
            "Could not read this DOCX file. Try re-saving it as .docx or upload a PDF."
        ) from exc
    text = _normalize_text("\n".join(parts))
    if not text:
        raise ValueError("This DOCX file contains no readable text.")
    return text


def _ocr_fallback(file_bytes: bytes) -> str:
    """Rasterizes a scanned PDF and runs OCR. Best-effort: any failure returns
    empty text so the caller raises the standard no-text error.

    Uses pypdfium2 for rasterization because it ships a bundled binary (no
    system poppler/pdftoppm dependency, unlike pdf2image) — this keeps OCR
    working on minimal containers and bare-metal hosts alike.
    """
    try:
        import pytesseract
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(file_bytes)
        try:
            chunks: list[str] = []
            for page in pdf:
                bitmap = page.render(scale=2.0)  # 2x for better OCR accuracy
                pil_image = bitmap.to_pil()
                chunks.append(pytesseract.image_to_string(pil_image))
            return "\n".join(chunks).strip()
        finally:
            pdf.close()
    except Exception as exc:
        logger.warning("ocr_fallback_failed", error=str(exc))
        return ""


def validate_upload(file_bytes: bytes) -> None:
    if not file_bytes:
        raise ValueError("Uploaded file is empty.")
    if len(file_bytes) > MAX_FILE_BYTES:
        raise ValueError("File is too large (max 10 MB).")


def parse_resume(filename: str, file_bytes: bytes) -> str:
    validate_upload(file_bytes)
    lower_name = (filename or "").lower()
    if lower_name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    if lower_name.endswith(".docx"):
        return parse_docx(file_bytes)
    raise ValueError("Unsupported file type. Please upload a PDF or DOCX.")
